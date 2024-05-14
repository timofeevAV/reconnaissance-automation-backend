from rest_framework import status, generics, views
from .models import Trip, TripDate
from .serializers import TripSerializer, TripDateSerializer, TripDetailsSerializer
from rest_framework.pagination import PageNumberPagination
from rest_framework import filters
from rest_framework.response import Response
from django.shortcuts import get_object_or_404
from django.http import HttpResponse
from docx import Document
from docx.shared import Pt, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import RGBColor
from urllib.parse import quote
from django.db.models import Q
from rest_framework.permissions import AllowAny
from django.contrib.auth import get_user_model
from PIL import Image


class Pagination(PageNumberPagination):
    page_size = 10
    page_size_query_param = 'perPage'
    max_page_size = 25


class TripsListView(generics.ListCreateAPIView):
    serializer_class = TripSerializer
    pagination_class = Pagination
    filter_backends = [filters.SearchFilter, filters.OrderingFilter]
    search_fields = ['name']
    ordering_fields = ['name', 'createdAt', 'updatedAt']

    def get_queryset(self):
        user = self.request.user
        return Trip.objects.filter(Q(owner=user) | Q(editors=user)).distinct()

    def perform_create(self, serializer):
        serializer.save(owner=self.request.user)


class TripDetailsView(generics.RetrieveUpdateDestroyAPIView):
    queryset = Trip.objects.all()
    serializer_class = TripDetailsSerializer


class TripEditorView(views.APIView):
    def post(self, request, pk, action):
        trip = get_object_or_404(Trip, pk=pk)

        if request.user != trip.owner:
            return Response({'error': 'Only the owner can add editors'}, status=status.HTTP_403_FORBIDDEN)

        if action == 'add_editor':
            editor_id = request.data.get('editor_id')

            if editor_id == str(trip.owner.id):
                return Response({'error': 'Owner cannot be added as an editor'}, status=status.HTTP_400_BAD_REQUEST)

            editor_user = get_user_model().objects.get(pk=editor_id)
            trip.add_editor(editor_user)
            return Response({'status': 'Editor added successfully'}, status=status.HTTP_200_OK)

        elif action == 'remove_editor':
            editor_id = request.data.get('editor_id')
            editor_user = get_user_model().objects.get(pk=editor_id)
            trip.remove_editor(editor_user)
            return Response({'status': 'Editor removed successfully'}, status=status.HTTP_200_OK)

        return Response({'error': 'Invalid action'}, status=status.HTTP_400_BAD_REQUEST)


class TripDownloadView(views.APIView):
    permission_classes = [AllowAny]

    def get(self, request, pk):
        trip = get_object_or_404(Trip, pk=pk)

        doc = Document()

        self.add_dates_section(doc, trip)
        self.add_specialists_section(doc, trip)
        self.add_recon_scheme_section(doc, trip)
        self.add_recon_points_section(doc, trip)

        file_name = self.generate_file_name(trip)
        response = HttpResponse(
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = f'attachment; filename={file_name}'
        doc.save(response)
        return response

    def add_dates_section(self, doc, trip):
        title = doc.add_heading('Даты проведения:', level=1)
        self.set_run_properties(title.runs[0])
        doc.add_paragraph(
            ", ".join([date.day.strftime("%d.%m.%Y") for date in trip.dates.all()]))

    def add_specialists_section(self, doc, trip):
        title = doc.add_heading('Проведена специалистами:', level=1)
        self.set_run_properties(title.runs[0])
        table = doc.add_table(rows=1, cols=1)
        table.style = 'Table Grid'
        first_row = table.rows[0].cells
        first_row[0].text = f'{trip.owner.lastName} {trip.owner.firstName} {trip.owner.middleName or ""}'
        for editor in trip.editors.all():
            row_cells = table.add_row().cells
            row_cells[0].text = f'{editor.lastName} {editor.firstName} {trip.owner.middleName or ""}'
        doc.add_page_break()

    def calculate_fit_size(self, original_width, original_height, max_width, max_height):
        ratio_width = max_width / original_width
        ratio_height = max_height / original_height
        ratio = min(ratio_width, ratio_height)

        new_width = int(original_width * ratio)
        new_height = int(original_height * ratio)

        return new_width, new_height

    def add_recon_scheme_section(self, doc, trip):
        title = doc.add_heading('Cхема рекогносцировки', level=1)
        self.set_run_properties(title.runs[0])
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        section = doc.sections[0]

        if trip.scheme:
            img = Image.open(trip.scheme.path)
            original_width, original_height = img.size
            print(section.page_width)
            max_width = section.page_width - section.left_margin - section.right_margin
            max_height = section.page_height - section.top_margin - section.bottom_margin
            new_width, new_height = self.calculate_fit_size(
                original_width, original_height, max_width, max_height
            )
            doc.add_picture(trip.scheme, width=Emu(
                new_width), height=Emu(new_height))

        doc.add_page_break()

    def add_recon_points_section(self, doc, trip):
        table = doc.add_table(rows=1, cols=4)
        table.alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.autofit = True
        table.allow_autofit = True
        table.style = 'Table Grid'

        headers = table.rows[0].cells
        headers[0].text = 'Номер точки'
        headers[1].text = 'Координаты'
        headers[2].text = 'Описание'
        headers[3].text = 'Фото'
        self.make_rows_bold(table.rows[0])

        # Set column widths
        table.columns[0].width = Cm(1.7)
        table.columns[1].width = Cm(3)
        table.columns[2].width = Cm(5.5)
        table.columns[3].width = Cm(5.4)

        for number, point in enumerate(trip.points.all()):
            row_cells = table.add_row().cells
            row_cells[0].text = str(number + 1)
            row_cells[1].text = f'Широта: {point.latitude}, Долгота: {point.longitude}'
            row_cells[2].text = point.description or ''
            cell_width = Cm(5)  # Adjust the width as needed
            cell_height = Cm(5)  # Adjust the height as needed
            self.add_photos_to_cell(
                row_cells[3], point.photos.all(), cell_width, cell_height)

    def set_run_properties(self, run):
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 0, 0)

    def make_rows_bold(self, row):
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True

    def add_photos_to_cell(self, cell, photos, width, height):
        paragraph = cell.paragraphs[0]
        run = paragraph.add_run()

        for photo in photos:
            run.add_picture(photo.photo, width=width, height=height)

    def generate_file_name(self, trip):
        trip_name_encoded = quote(f'ЖР_{trip.name.replace(" ", "_")}')
        return f'{trip_name_encoded}_{trip.id}.docx'


class DeleteTripsView(generics.DestroyAPIView):
    queryset = Trip.objects.all()
    serializer_class = TripSerializer

    def delete(self, request, *args, **kwargs):
        trip_ids = request.data.get('trip_ids', [])

        queryset = self.get_queryset().filter(id__in=trip_ids)

        if not queryset.exists():
            return Response({'detail': 'Объекты с указанными идентификаторами не найдены.'}, status=status.HTTP_404_NOT_FOUND)

        queryset.delete()

        return Response({'detail': 'Объекты успешно удалены.'}, status=status.HTTP_204_NO_CONTENT)


class TripDateListView(generics.ListCreateAPIView):
    serializer_class = TripDateSerializer

    def get_queryset(self):
        trip = get_object_or_404(Trip, pk=self.kwargs['trip_id'])
        return trip.dates.all()

    def post(self, request, trip_id, format=None):
        day = request.data.get('day')
        if day:
            try:
                trip = Trip.objects.get(pk=trip_id)
            except Trip.DoesNotExist:
                return Response({'detail': 'Нет выезда с таким id'}, status=status.HTTP_400_BAD_REQUEST)
            trip_date = TripDate.objects.create(day=day, trip=trip)
            serialized_trip_date = TripDateSerializer(trip_date)
            return Response(serialized_trip_date.data, status=status.HTTP_201_CREATED)
        else:
            return Response(status=status.HTTP_400_BAD_REQUEST)


class TripDateDetailsView(generics.DestroyAPIView):
    queryset = TripDate.objects.all()
    serializer_class = TripDateSerializer
